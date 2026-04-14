"""
Boundary Decision Prediction — Circular boundary classification
with ordinal probability estimation.

Classifies data points (I, F) into 4 ordered exposure classes (E)
separated by concentric circle boundaries centred on a user-defined
point (default: centroid of class 1).

Step 1: Optimal radii found by minimising classification error.
Step 2: Cumulative ordinal logistic regression on
        r = sqrt((I-I0)²+(F-F0)²) yields calibrated class
        probabilities P(E=k | I, F).
"""

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
from scipy.optimize import minimize
from itertools import combinations
from pathlib import Path

# ---------------------------------------------------------------------------
# 1. Synthetic dataset generation
# ---------------------------------------------------------------------------

def generate_dataset(n_total: int = 54, seed: int = 42,
                     full_plane: bool = True) -> pd.DataFrame:
    """
    Generate *n_total* points with 4 classes (E = 1..4) separated
    by three concentric circles of true radii r1 < r2 < r3.

    If *full_plane* is True the points are spread over all quadrants
    (θ ∈ [0, 2π]); otherwise they are restricted to the first quadrant.
    """
    rng = np.random.default_rng(seed)

    true_radii = [3.0, 6.0, 9.0]
    n_per_class = n_total // 4
    remainder = n_total - 4 * n_per_class

    theta_max = 2 * np.pi if full_plane else np.pi / 2

    records = []
    bounds = [0.0] + true_radii + [13.0]

    for cls_idx in range(4):
        r_lo, r_hi = bounds[cls_idx], bounds[cls_idx + 1]
        n_cls = n_per_class + (1 if cls_idx < remainder else 0)

        r = rng.uniform(r_lo + 0.3, r_hi - 0.3, size=n_cls)
        theta = rng.uniform(0, theta_max, size=n_cls)

        I_vals = r * np.cos(theta)
        F_vals = r * np.sin(theta)

        for i in range(n_cls):
            records.append({"I": I_vals[i], "F": F_vals[i], "E": cls_idx + 1})

    df = pd.DataFrame(records)
    return df.sample(frac=1, random_state=seed).reset_index(drop=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _dist_to_center(I: np.ndarray, F: np.ndarray,
                    center: tuple[float, float]) -> np.ndarray:
    """Euclidean distance from each point to *center*."""
    return np.sqrt((I - center[0])**2 + (F - center[1])**2)


def compute_default_center(df: pd.DataFrame) -> tuple[float, float]:
    """Return the centroid of class E = 1 (lowest exposure)."""
    c1 = df[df["E"] == df["E"].min()]
    return float(c1["I"].mean()), float(c1["F"].mean())


# ---------------------------------------------------------------------------
# 2. Optimal circle-boundary fitting
# ---------------------------------------------------------------------------

def classify_by_radii(I: np.ndarray, F: np.ndarray,
                      radii: np.ndarray,
                      center: tuple[float, float] = (0.0, 0.0)) -> np.ndarray:
    """Assign class labels 1..4 based on distance to *center* and 3 radii."""
    dist = _dist_to_center(I, F, center)
    labels = np.ones(len(dist), dtype=int)
    for k, r in enumerate(sorted(radii)):
        labels[dist > r] = k + 2
    return labels


def fit_circle_boundaries(df: pd.DataFrame,
                          center: tuple[float, float] = (0.0, 0.0),
                          n_classes: int = 4):
    """
    Find the (n_classes-1) radii that best separate the classes
    around *center* by minimising the misclassification rate.
    """
    I = df["I"].values
    F = df["F"].values
    y_true = df["E"].values

    r_max = _dist_to_center(I, F, center).max()

    def objective(params):
        radii = np.sort(params)
        y_pred = classify_by_radii(I, F, radii, center)
        error = np.mean(y_pred != y_true)
        penalty = 1e3 * np.sum(np.maximum(0, -np.diff(radii)))
        return error + penalty

    n_boundaries = n_classes - 1
    grid = np.linspace(0.5, r_max * 0.95, 6 * n_boundaries)

    best_result = None
    for combo in combinations(grid, n_boundaries):
        combo = sorted(combo)
        if any(combo[i] >= combo[i + 1] for i in range(len(combo) - 1)):
            continue
        res = minimize(objective, list(combo), method="Nelder-Mead",
                       options={"maxiter": 5000, "xatol": 1e-6})
        if best_result is None or res.fun < best_result.fun:
            best_result = res

    optimal_radii = np.sort(best_result.x)
    return optimal_radii, best_result.fun


# ---------------------------------------------------------------------------
# 3. Publication-quality plot — boundaries
# ---------------------------------------------------------------------------

CLASS_LABELS = {1: "E = 1", 2: "E = 2", 3: "E = 3", 4: "E = 4"}
CLASS_COLORS = {1: "#1b9e77", 2: "#d95f02", 3: "#7570b3", 4: "#e7298a"}
CLASS_MARKERS = {1: "o", 2: "s", 3: "^", 4: "D"}


def _data_extent(df: pd.DataFrame, radii: np.ndarray,
                 center: tuple[float, float], margin: float = 1.5):
    """Compute the outer radius that encompasses data + outer boundary."""
    r_outer = max(
        radii.max(),
        _dist_to_center(df["I"].values, df["F"].values, center).max()
    ) + margin
    return r_outer


def plot_boundaries(df: pd.DataFrame, radii: np.ndarray,
                    accuracy: float,
                    center: tuple[float, float] = (0.0, 0.0),
                    save_path: str = "boundary_plot.png"):
    """Scientific-quality 2D scatter with full-circle boundaries."""

    plt.rcParams.update({
        "font.family": "serif", "font.size": 11,
        "axes.linewidth": 0.8,
        "xtick.direction": "in", "ytick.direction": "in",
        "xtick.major.size": 4, "ytick.major.size": 4,
        "figure.dpi": 150,
    })

    cx, cy = center
    fig, ax = plt.subplots(figsize=(7.5, 7))

    for cls in sorted(df["E"].unique()):
        subset = df[df["E"] == cls]
        ax.scatter(subset["I"], subset["F"],
                   c=CLASS_COLORS[cls], marker=CLASS_MARKERS[cls],
                   s=60, edgecolors="k", linewidths=0.5,
                   label=CLASS_LABELS[cls], zorder=5)

    theta_full = np.linspace(0, 2 * np.pi, 500)
    r_outer = _data_extent(df, radii, center)
    all_bounds = [0.0] + list(radii) + [r_outer]

    ax.fill(cx + all_bounds[1] * np.cos(theta_full),
            cy + all_bounds[1] * np.sin(theta_full),
            color=CLASS_COLORS[1], alpha=0.10, zorder=1)
    for i in range(1, 4):
        r_in, r_out = all_bounds[i], all_bounds[i + 1]
        x_o = cx + r_out * np.cos(theta_full)
        y_o = cy + r_out * np.sin(theta_full)
        x_i = cx + r_in * np.cos(theta_full[::-1])
        y_i = cy + r_in * np.sin(theta_full[::-1])
        ax.fill(np.concatenate([x_o, x_i]),
                np.concatenate([y_o, y_i]),
                color=CLASS_COLORS[i + 1], alpha=0.10, zorder=1)

    for k, r in enumerate(radii):
        x_c = cx + r * np.cos(theta_full)
        y_c = cy + r * np.sin(theta_full)
        ax.plot(x_c, y_c, color="k", linewidth=1.2, linestyle="--", zorder=4)
        ax.annotate(
            f"$r_{k+1} = {r:.2f}$",
            xy=(cx + r * np.cos(np.pi / 4),
                cy + r * np.sin(np.pi / 4)),
            xytext=(12, 12), textcoords="offset points",
            fontsize=9, fontstyle="italic",
            arrowprops=dict(arrowstyle="->", color="grey", lw=0.8),
            bbox=dict(boxstyle="round,pad=0.25", fc="white",
                      ec="grey", alpha=0.9),
            zorder=6)

    ax.plot(cx, cy, marker="+", ms=14, mew=2.5, color="red", zorder=8)
    ax.annotate(f"centre ({cx:.2f}, {cy:.2f})",
                xy=(cx, cy), xytext=(14, -14),
                textcoords="offset points", fontsize=8.5,
                color="red", fontweight="bold", zorder=8)

    all_I = df["I"].values
    all_F = df["F"].values
    pad = 1.0
    xlim_lo = min(all_I.min(), cx - r_outer) - pad
    xlim_hi = max(all_I.max(), cx + r_outer) + pad
    ylim_lo = min(all_F.min(), cy - r_outer) - pad
    ylim_hi = max(all_F.max(), cy + r_outer) + pad
    ax.set_xlim(xlim_lo, xlim_hi)
    ax.set_ylim(ylim_lo, ylim_hi)

    ax.axhline(0, color="k", lw=0.4, zorder=2)
    ax.axvline(0, color="k", lw=0.4, zorder=2)
    ax.set_xlabel("$I$", fontsize=13)
    ax.set_ylabel("$F$", fontsize=13)
    ax.set_title("Boundary Decision Prediction\n"
                 "Circular classification boundaries",
                 fontsize=13, fontweight="bold", pad=12)
    ax.set_aspect("equal")
    ax.legend(loc="upper right", framealpha=0.9, edgecolor="grey",
              fontsize=10, title="Classes", title_fontsize=11)

    acc_text = f"Accuracy = {accuracy:.1%}"
    eq_lines = "\n".join(
        [f"$r_{k+1}$: $(I-I_0)^2+(F-F_0)^2 = {r**2:.2f}$"
         for k, r in enumerate(radii)])
    center_text = f"Centre $(I_0, F_0) = ({cx:.2f},\\;{cy:.2f})$"
    info_text = f"{acc_text}\n{center_text}\n\nBoundary equations:\n{eq_lines}"
    ax.text(0.02, 0.98, info_text, transform=ax.transAxes,
            fontsize=8.5, verticalalignment="top",
            bbox=dict(boxstyle="round,pad=0.5", fc="lightyellow",
                      ec="grey", alpha=0.9), zorder=7)

    ax.grid(True, linestyle=":", linewidth=0.4, alpha=0.6)
    fig.tight_layout()
    fig.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    print(f"Plot saved to {save_path}")
    return save_path


# ---------------------------------------------------------------------------
# 4. Cumulative ordinal logistic regression
# ---------------------------------------------------------------------------

def _sigmoid(x):
    return 1.0 / (1.0 + np.exp(-np.clip(x, -500, 500)))


def ordinal_cumulative_probs(r: np.ndarray,
                             alpha: np.ndarray,
                             beta: float) -> np.ndarray:
    """
    Compute P(E <= k | r) for k = 1..K-1 using the cumulative model.
    Returns shape (n_samples, K) with class probabilities P(E = k | r).
    """
    K = len(alpha) + 1
    n = len(r)
    cum = np.zeros((n, K - 1))
    for k in range(K - 1):
        cum[:, k] = _sigmoid(alpha[k] - beta * r)

    probs = np.zeros((n, K))
    probs[:, 0] = cum[:, 0]
    for k in range(1, K - 1):
        probs[:, k] = cum[:, k] - cum[:, k - 1]
    probs[:, K - 1] = 1.0 - cum[:, K - 2]
    probs = np.clip(probs, 1e-12, 1.0)
    probs /= probs.sum(axis=1, keepdims=True)
    return probs


def ordinal_neg_log_likelihood(params, r, y, K=4):
    """Negative log-likelihood for the cumulative ordinal model."""
    alpha = params[:K - 1]
    beta = params[K - 1]
    probs = ordinal_cumulative_probs(r, alpha, beta)
    nll = 0.0
    for i in range(len(r)):
        nll -= np.log(probs[i, y[i] - 1])
    return nll


def fit_ordinal_logistic(df: pd.DataFrame, radii: np.ndarray,
                         center: tuple[float, float] = (0.0, 0.0)):
    """
    Fit a cumulative ordinal logistic model on
    r = sqrt((I-I0)² + (F-F0)²).
    """
    r = _dist_to_center(df["I"].values, df["F"].values, center)
    y = df["E"].values
    K = 4

    def objective(params):
        alpha = params[:K - 1]
        beta = params[K - 1]
        if beta <= 0:
            return 1e12
        if not np.all(np.diff(alpha) > 0):
            return 1e12
        return ordinal_neg_log_likelihood(params, r, y, K)

    best_result = None
    for beta_try in [0.5, 1.0, 2.0, 3.0, 5.0]:
        alpha_try = beta_try * radii
        x0 = np.concatenate([alpha_try, [beta_try]])
        res = minimize(objective, x0, method="Nelder-Mead",
                       options={"maxiter": 20000, "xatol": 1e-10,
                                "fatol": 1e-10})
        if best_result is None or res.fun < best_result.fun:
            best_result = res

    alpha_opt = best_result.x[:K - 1]
    beta_opt = best_result.x[K - 1]

    probs = ordinal_cumulative_probs(r, alpha_opt, beta_opt)
    y_pred = probs.argmax(axis=1) + 1
    accuracy = np.mean(y_pred == y)

    return alpha_opt, beta_opt, probs, accuracy


# ---------------------------------------------------------------------------
# 5. Probability visualisation (publication-quality)
# ---------------------------------------------------------------------------

EXPOSURE_LABELS = {
    1: "E=1 (faible)",
    2: "E=2 (modérée)",
    3: "E=3 (élevée)",
    4: "E=4 (très élevée)",
}


def plot_ordinal_probabilities(df: pd.DataFrame, radii: np.ndarray,
                               alpha: np.ndarray, beta: float,
                               accuracy: float,
                               center: tuple[float, float] = (0.0, 0.0),
                               save_path: str = "probability_plot.png"):
    """
    Two-panel figure:
      Left  – 2D scatter with colour = true class, size ∝ confidence.
      Right – P(E=k | r) curves as a function of r.
    """
    plt.rcParams.update({
        "font.family": "serif", "font.size": 11,
        "axes.linewidth": 0.8,
        "xtick.direction": "in", "ytick.direction": "in",
        "xtick.major.size": 4, "ytick.major.size": 4,
        "figure.dpi": 150,
    })

    cx, cy = center
    r_data = _dist_to_center(df["I"].values, df["F"].values, center)
    probs_data = ordinal_cumulative_probs(r_data, alpha, beta)
    max_prob = probs_data.max(axis=1)

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14.5, 6.5),
                                    gridspec_kw={"width_ratios": [1, 1.05]})

    # --- Left panel ---
    theta_full = np.linspace(0, 2 * np.pi, 500)
    r_outer = _data_extent(df, radii, center)
    all_bounds = [0.0] + list(radii) + [r_outer]

    ax1.fill(cx + all_bounds[1] * np.cos(theta_full),
             cy + all_bounds[1] * np.sin(theta_full),
             color=CLASS_COLORS[1], alpha=0.08, zorder=1)
    for i in range(1, 4):
        r_in, r_out = all_bounds[i], all_bounds[i + 1]
        x_o = cx + r_out * np.cos(theta_full)
        y_o = cy + r_out * np.sin(theta_full)
        x_i = cx + r_in * np.cos(theta_full[::-1])
        y_i = cy + r_in * np.sin(theta_full[::-1])
        ax1.fill(np.concatenate([x_o, x_i]),
                 np.concatenate([y_o, y_i]),
                 color=CLASS_COLORS[i + 1], alpha=0.08, zorder=1)

    for k, rv in enumerate(radii):
        ax1.plot(cx + rv * np.cos(theta_full),
                 cy + rv * np.sin(theta_full),
                 color="k", lw=1.0, ls="--", alpha=0.5, zorder=3)

    sizes = 30 + 120 * (max_prob - max_prob.min()) / \
        (max_prob.max() - max_prob.min() + 1e-9)
    for cls in sorted(df["E"].unique()):
        mask = df["E"].values == cls
        ax1.scatter(df["I"].values[mask], df["F"].values[mask],
                    c=CLASS_COLORS[cls], marker=CLASS_MARKERS[cls],
                    s=sizes[mask], edgecolors="k", linewidths=0.4,
                    label=EXPOSURE_LABELS[cls], zorder=5)

    ax1.plot(cx, cy, marker="+", ms=12, mew=2, color="red", zorder=8)

    all_I = df["I"].values
    all_F = df["F"].values
    pad = 1.0
    ax1.set_xlim(min(all_I.min(), cx - r_outer) - pad,
                 max(all_I.max(), cx + r_outer) + pad)
    ax1.set_ylim(min(all_F.min(), cy - r_outer) - pad,
                 max(all_F.max(), cy + r_outer) + pad)
    ax1.axhline(0, color="k", lw=0.4, zorder=2)
    ax1.axvline(0, color="k", lw=0.4, zorder=2)
    ax1.set_xlabel("$I$", fontsize=13)
    ax1.set_ylabel("$F$", fontsize=13)
    ax1.set_title("Classification avec confiance\n"
                  "(taille $\\propto$ max $P(E=k|r)$)",
                  fontsize=12, fontweight="bold", pad=10)
    ax1.set_aspect("equal")
    ax1.legend(loc="upper right", fontsize=8.5, framealpha=0.9,
               edgecolor="grey", title="Exposition", title_fontsize=9.5)
    ax1.grid(True, ls=":", lw=0.4, alpha=0.5)

    # --- Right panel ---
    r_max_plot = r_data.max() * 1.3
    r_smooth = np.linspace(0, r_max_plot, 500)
    probs_smooth = ordinal_cumulative_probs(r_smooth, alpha, beta)

    line_styles = ["-", "--", "-.", ":"]
    for k in range(4):
        ax2.fill_between(r_smooth, probs_smooth[:, k],
                         color=CLASS_COLORS[k + 1], alpha=0.18, zorder=1)
        ax2.plot(r_smooth, probs_smooth[:, k],
                 color=CLASS_COLORS[k + 1], lw=2.2, ls=line_styles[k],
                 label=EXPOSURE_LABELS[k + 1], zorder=2)

    for rv in radii:
        ax2.axvline(rv, color="grey", lw=0.8, ls=":", alpha=0.6)

    for i, row in df.iterrows():
        ri = np.sqrt((row["I"] - cx)**2 + (row["F"] - cy)**2)
        cls = int(row["E"])
        prob_k = probs_data[i, cls - 1]
        ax2.scatter(ri, prob_k, c=CLASS_COLORS[cls],
                    marker=CLASS_MARKERS[cls],
                    s=30, edgecolors="k", linewidths=0.3,
                    zorder=5, alpha=0.7)

    ax2.set_xlabel(
        "$r = \\sqrt{(I-I_0)^2 + (F-F_0)^2}$", fontsize=13)
    ax2.set_ylabel("$P(E = k \\mid r)$", fontsize=13)
    ax2.set_title("Probabilités ordinales cumulatives\n"
                  "Régression logistique ordinale",
                  fontsize=12, fontweight="bold", pad=10)
    ax2.set_xlim(0, r_max_plot)
    ax2.set_ylim(-0.02, 1.05)
    ax2.legend(loc="center right", fontsize=8.5, framealpha=0.9,
               edgecolor="grey", title="Exposition", title_fontsize=9.5)
    ax2.grid(True, ls=":", lw=0.4, alpha=0.5)

    eq_text = (
        f"$P(E \\leq k \\mid r) = \\sigma(\\alpha_k - \\beta \\, r)$\n"
        f"$r = \\sqrt{{(I-{cx:.2f})^2 + (F-{cy:.2f})^2}}$\n\n"
        f"$\\alpha_1 = {alpha[0]:.3f}$\n"
        f"$\\alpha_2 = {alpha[1]:.3f}$\n"
        f"$\\alpha_3 = {alpha[2]:.3f}$\n"
        f"$\\beta = {beta:.3f}$\n\n"
        f"Précision = {accuracy:.1%}")
    ax2.text(0.98, 0.98, eq_text, transform=ax2.transAxes,
             fontsize=8.5, verticalalignment="top",
             horizontalalignment="right",
             bbox=dict(boxstyle="round,pad=0.5", fc="lightyellow",
                       ec="grey", alpha=0.9), zorder=7)

    fig.tight_layout(w_pad=3)
    fig.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    print(f"Probability plot saved to {save_path}")
    return save_path


# ---------------------------------------------------------------------------
# 6. DOCX report generation
# ---------------------------------------------------------------------------

def generate_report(radii, alpha, beta, accuracy_hard, accuracy_ordinal,
                    center: tuple[float, float] = (0.0, 0.0),
                    save_path="rapport_methode.docx"):
    """Generate a short scientific report in DOCX format."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cx, cy = center

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Boundary Decision Prediction", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Classification par frontières circulaires "
        "et estimation de probabilités ordinales"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_heading("1. Contexte", level=1)
    doc.add_paragraph(
        "On considère un problème de classification de données "
        "bidimensionnelles (I, F) en 4 classes ordonnées d'exposition "
        "E ∈ {1, 2, 3, 4}, où E = 1 correspond à une faible exposition "
        "énergétique et E = 4 à une exposition très élevée. "
        "Les frontières de séparation sont des cercles concentriques "
        f"centrés en un point (I₀, F₀) = ({cx:.4f}, {cy:.4f}), "
        "choisi par défaut comme le barycentre de la classe 1 "
        "(exposition la plus faible)."
    )

    doc.add_heading("2. Étape 1 — Détermination des frontières", level=1)
    doc.add_paragraph(
        "La variable d'exposition combinée est définie comme la distance "
        "au centre :"
    )
    doc.add_paragraph(
        f"    r = √((I − I₀)² + (F − F₀)²)    "
        f"avec (I₀, F₀) = ({cx:.4f}, {cy:.4f})",
        style="No Spacing")
    doc.add_paragraph(
        "Les 4 classes sont séparées par 3 cercles concentriques de rayons "
        "r₁ < r₂ < r₃. Les rayons optimaux sont déterminés par minimisation "
        "du taux de mauvaise classification via Nelder-Mead (scipy.optimize)."
    )
    doc.add_paragraph("Équations des frontières :")
    for k, r in enumerate(radii):
        doc.add_paragraph(
            f"    Frontière {k+1} :  "
            f"(I − I₀)² + (F − F₀)² = {r**2:.4f}   "
            f"(r{k+1} = {r:.4f})",
            style="No Spacing")
    doc.add_paragraph(
        f"\nPrécision de classification dure : {accuracy_hard:.1%}")

    doc.add_heading(
        "3. Étape 2 — Régression logistique ordinale cumulative", level=1)
    doc.add_paragraph(
        "Pour obtenir des probabilités d'appartenance à chaque classe, on "
        "utilise un modèle de régression logistique ordinale cumulative "
        "(proportional odds model). Ce modèle est adapté car :"
    )
    for b in [
        "Les classes ont un ordre naturel (intensité d'exposition "
        "croissante).",
        "La variable latente sous-jacente (l'exposition r) est continue.",
        "Le modèle garantit la cohérence des probabilités avec "
        "l'ordinalité.",
        "La formulation est invariante par rotation autour du centre "
        "choisi.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.1. Formulation mathématique", level=2)
    doc.add_paragraph("Le modèle cumulatif s'écrit :")
    doc.add_paragraph(
        "    P(E ≤ k | r) = σ(αₖ − β · r),    k = 1, 2, 3",
        style="No Spacing")
    doc.add_paragraph(
        "où σ(x) = 1 / (1 + exp(−x)) est la fonction sigmoïde logistique, "
        "α₁ < α₂ < α₃ sont les seuils et β > 0 est la pente commune. "
        "La variable explicative est r = √((I−I₀)² + (F−F₀)²)."
    )
    doc.add_paragraph(
        "Les probabilités par classe se déduisent par différence :")
    for eq in [
        "P(E = 1 | r) = P(E ≤ 1 | r)",
        "P(E = 2 | r) = P(E ≤ 2 | r) − P(E ≤ 1 | r)",
        "P(E = 3 | r) = P(E ≤ 3 | r) − P(E ≤ 2 | r)",
        "P(E = 4 | r) = 1 − P(E ≤ 3 | r)",
    ]:
        doc.add_paragraph(f"    {eq}", style="No Spacing")

    doc.add_heading("3.2. Estimation des paramètres", level=2)
    doc.add_paragraph(
        "Les paramètres sont estimés par maximum de vraisemblance. "
        "Les valeurs initiales des seuils sont dérivées des rayons "
        "optimaux de l'étape 1 (αₖ⁰ = β⁰ · rₖ)."
    )
    doc.add_paragraph("Paramètres estimés :")
    for k in range(len(alpha)):
        doc.add_paragraph(
            f"    α{k+1} = {alpha[k]:.4f}", style="No Spacing")
    doc.add_paragraph(f"    β  = {beta:.4f}", style="No Spacing")
    doc.add_paragraph(
        f"\nPrécision par probabilité maximale : {accuracy_ordinal:.1%}")

    doc.add_heading("3.3. Interprétation physique", level=2)
    doc.add_paragraph(
        "Les seuils αₖ/β correspondent aux rayons de transition. "
        "La pente β contrôle la netteté : plus β est grand, plus la "
        "transition est abrupte. Aux frontières (r ≈ αₖ/β), l'incertitude "
        "est maximale."
    )
    r_transitions = alpha / beta
    for k in range(len(alpha)):
        doc.add_paragraph(
            f"    Transition E={k+1}/E={k+2} : "
            f"r = α{k+1}/β = {r_transitions[k]:.4f}",
            style="No Spacing")

    doc.add_heading("3.4. Choix du centre", level=2)
    doc.add_paragraph(
        "Le centre des cercles (I₀, F₀) est configurable par l'utilisateur. "
        "Par défaut, il est fixé au barycentre de la classe E = 1 "
        "(exposition la plus faible), ce qui est physiquement cohérent : "
        "le point de moindre exposition constitue un centre naturel "
        "autour duquel l'exposition croît radialement."
    )

    doc.add_heading("4. Conclusion", level=1)
    doc.add_paragraph(
        "La combinaison d'une classification par frontières circulaires "
        "centrées et d'une régression logistique ordinale fournit un cadre "
        "complet pour :\n"
        "  (a) classifier les points en classes d'exposition,\n"
        "  (b) quantifier l'incertitude via des probabilités calibrées,\n"
        "  (c) identifier les zones de transition ambiguës.\n\n"
        "Le centre des cercles est paramétrable, permettant d'adapter "
        "le modèle à différentes configurations géométriques des données."
    )

    doc.save(save_path)
    print(f"Report saved to {save_path}")
    return save_path


# ---------------------------------------------------------------------------
# 7. Main
# ---------------------------------------------------------------------------

def main():
    print("=" * 60)
    print("  BOUNDARY DECISION PREDICTION")
    print("  Circular classification + ordinal probabilities")
    print("=" * 60)

    df = generate_dataset(n_total=54, seed=42, full_plane=True)
    csv_path = Path("dataset.csv")
    df.to_csv(csv_path, index=False)
    print(f"\nDataset ({len(df)} samples) saved to {csv_path}")
    print(df.groupby("E").size().rename("count").to_frame())

    # Centre = centroid of class 1 (default) — override here if needed
    center = compute_default_center(df)
    print(f"\nCircle centre (centroid of E=1): "
          f"I₀ = {center[0]:.4f}, F₀ = {center[1]:.4f}")

    print("\nFitting optimal circle boundaries ...")
    radii, error = fit_circle_boundaries(df, center=center)
    accuracy_hard = 1.0 - error

    print(f"\nOptimal radii found:")
    for k, r in enumerate(radii):
        print(f"  r{k+1} = {r:.4f}  →  (I-I₀)²+(F-F₀)² = {r**2:.4f}")
    print(f"\nHard classification accuracy: {accuracy_hard:.1%}")

    y_pred = classify_by_radii(df["I"].values, df["F"].values,
                               radii, center)
    df["E_pred"] = y_pred

    print("\nFitting cumulative ordinal logistic model ...")
    alpha, beta, probs, accuracy_ordinal = fit_ordinal_logistic(
        df, radii, center=center)

    print(f"\nOrdinal model parameters:")
    for k in range(len(alpha)):
        print(f"  α{k+1} = {alpha[k]:.4f}")
    print(f"  β  = {beta:.4f}")
    print(f"\nTransition radii (αk/β):")
    for k in range(len(alpha)):
        print(f"  r{k+1}* = {alpha[k]/beta:.4f}")
    print(f"\nOrdinal classification accuracy: {accuracy_ordinal:.1%}")

    df["P(E=1)"] = probs[:, 0]
    df["P(E=2)"] = probs[:, 1]
    df["P(E=3)"] = probs[:, 2]
    df["P(E=4)"] = probs[:, 3]
    df["E_ordinal"] = probs.argmax(axis=1) + 1
    df.to_csv("dataset_with_probabilities.csv", index=False)
    print("\nProbability vectors (first 10 rows):")
    print(df[["I", "F", "E", "P(E=1)", "P(E=2)", "P(E=3)", "P(E=4)"]].head(10)
          .to_string(index=False, float_format="%.4f"))

    plot_boundaries(df, radii, accuracy_hard, center=center)
    plot_ordinal_probabilities(df, radii, alpha, beta,
                               accuracy_ordinal, center=center)
    generate_report(radii, alpha, beta, accuracy_hard, accuracy_ordinal,
                    center=center)

    print("\n" + "=" * 60)
    print("  All outputs generated:")
    print(f"    - {csv_path}")
    print(f"    - dataset_with_probabilities.csv")
    print(f"    - boundary_plot.png")
    print(f"    - probability_plot.png")
    print(f"    - rapport_methode.docx")
    print("=" * 60)


if __name__ == "__main__":
    main()
